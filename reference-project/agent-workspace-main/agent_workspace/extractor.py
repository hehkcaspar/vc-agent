"""Multi-format file extraction: text, docx, pdf, excel, images."""

from __future__ import annotations

import base64
import csv as csv_mod
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from .config import ExtractionConfig

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Encoding fallback chain
# ---------------------------------------------------------------------------

_ENCODINGS = ("utf-8", "gbk", "latin-1")


def _read_text_with_fallback(path: Path) -> str:
    for enc in _ENCODINGS:
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return f"[Unable to decode: {path.name}]"


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------

def _extract_text(path: Path, max_chars: int) -> str:
    text = _read_text_with_fallback(path)
    return text[:max_chars]


def _extract_docx(path: Path, max_chars: int) -> str:
    from docx import Document

    doc = Document(path)
    lines: List[str] = []
    char_count = 0

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
            char_count += len(t)
            if char_count >= max_chars:
                break

    if char_count < max_chars:
        for table in doc.tables:
            for row in table.rows:
                line = "\t".join(cell.text.strip() for cell in row.cells)
                lines.append(line)
                char_count += len(line)
                if char_count >= max_chars:
                    break
            if char_count >= max_chars:
                break

    return "\n".join(lines)[:max_chars]


def _extract_pdf(path: Path, max_chars: int) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    parts: List[str] = []
    total = 0
    for page in doc:
        text = page.get_text()
        parts.append(text)
        total += len(text)
        if total >= max_chars:
            break
    doc.close()
    return "\n".join(parts)[:max_chars]


def _extract_excel(path: Path, max_chars: int, max_sheets: int, max_rows: int) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: List[str] = []
    total = 0

    for ws in wb.worksheets[:max_sheets]:
        parts.append(f"[Sheet: {ws.title}]")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                parts.append("...(truncated)")
                break
            line = "\t".join("" if c is None else str(c) for c in row)
            parts.append(line)
            total += len(line)
            if total >= max_chars:
                break
        if total >= max_chars:
            break

    wb.close()
    return "\n".join(parts)[:max_chars]


def _extract_image(path: Path) -> Dict[str, str]:
    """Return base64 data URL for multimodal LLM consumption."""
    mime_map = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".gif": "image/gif",
        ".bmp": "image/bmp", ".webp": "image/webp",
    }
    mime = mime_map.get(path.suffix.lower(), "image/jpeg")
    data = base64.b64encode(path.read_bytes()).decode("utf-8")
    return {"mime": mime, "data_url": f"data:{mime};base64,{data}"}


def _extract_csv(path: Path, max_chars: int) -> str:
    """Extract CSV file content using stdlib csv module."""
    text = _read_text_with_fallback(path)
    parts: List[str] = []
    total = 0
    try:
        reader = csv_mod.reader(text.splitlines())
        for i, row in enumerate(reader):
            line = "\t".join(row)
            parts.append(line)
            total += len(line)
            if total >= max_chars:
                parts.append("...(truncated)")
                break
    except csv_mod.Error:
        # Fallback: treat as plain text
        return text[:max_chars]
    return "\n".join(parts)[:max_chars]


# ---------------------------------------------------------------------------
# Unified extraction entry point
# ---------------------------------------------------------------------------

def extract_file(
    path: Path,
    file_type: str,
    config: Optional[ExtractionConfig] = None,
) -> Dict[str, Any]:
    """Extract content from a single file.

    Returns: {type, path, text} or {type, path, image_url} for images.
    """
    cfg = config or ExtractionConfig()
    rel = str(path)

    try:
        if file_type == "text":
            return {"type": file_type, "path": rel, "text": _extract_text(path, cfg.max_text_chars)}

        if file_type == "word":
            return {"type": file_type, "path": rel, "text": _extract_docx(path, cfg.max_text_chars)}

        if file_type == "word_legacy":
            return {"type": file_type, "path": rel, "text": f"[Legacy .doc format: {path.name} — convert to .docx for full extraction]"}

        if file_type == "pdf":
            return {"type": file_type, "path": rel, "text": _extract_pdf(path, cfg.max_text_chars)}

        if file_type == "excel":
            text = _extract_excel(path, cfg.max_text_chars, cfg.max_excel_sheets, cfg.max_excel_rows)
            return {"type": file_type, "path": rel, "text": text}

        if file_type == "csv":
            return {"type": file_type, "path": rel, "text": _extract_csv(path, cfg.max_text_chars)}

        if file_type == "image":
            img = _extract_image(path)
            return {"type": file_type, "path": rel, "image_url": img["data_url"], "text": f"[Image: {path.name}]"}

        return {"type": file_type, "path": rel, "text": f"[Unsupported file type: {path.suffix}]"}

    except Exception as exc:
        logger.warning("Extraction failed for %s: %s", path, exc)
        return {"type": "error", "path": rel, "text": f"[Extraction failed: {exc}]"}
